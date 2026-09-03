"""@brief Viewer for Word documents: .docx via python-docx, .odt via raw XML.

@details DOCX content (headings, paragraphs with bold/italic/underline runs,
bullet lists, and tables) is converted to HTML and rendered in QTextBrowser.
ODT files are read directly from the OpenDocument ZIP archive (content.xml).
Everything is parsed in-process; no external office software is launched.
"""

import html
import os
import zipfile
import xml.etree.ElementTree as ET
from itertools import islice

from PyQt6.QtWidgets import QTextBrowser, QVBoxLayout, QWidget

from universal_viewer.viewers.base import BaseViewer

#: @brief OpenDocument text namespace prefix used when parsing content.xml.
_ODT_TEXT_NS = "{urn:oasis:names:tc:opendocument:xmlns:text:1.0}"

#: @brief Maximum number of table rows rendered (guards against huge sheets).
_MAX_TABLE_ROWS = 500

#: @brief Maximum number of content blocks (paragraphs/tables) rendered.
_MAX_BLOCKS = 50_000

#: @brief Maximum uncompressed size of content.xml inside an ODT file (64 MB).
_MAX_XML_BYTES = 64 * 1024 * 1024

#: @brief Maximum size of an office document opened by this viewer (256 MB).
_MAX_DOC_BYTES = 256 * 1024 * 1024

#: @brief Extensions handled by the document viewer.
_DOCUMENT_EXTENSIONS = frozenset({".docx", ".odt"})


class DocxViewer(BaseViewer):
    """@brief Rendered viewer for DOCX and ODT word-processing documents."""

    SUPPORTED_EXTENSIONS = _DOCUMENT_EXTENSIONS

    def __init__(self, parent: QWidget | None = None) -> None:
        """@brief Construct the widget with its browser area.

        @param parent: Optional parent widget.
        """
        super().__init__(parent)
        self._browser = QTextBrowser(self)
        self._browser.setOpenLinks(False)
        self._browser.setOpenExternalLinks(False)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self._browser)

    def load(self, path: str) -> None:
        """@brief Render a DOCX or ODT document as formatted HTML.

        @param path: Absolute path of the document to display.
        """
        extension = os.path.splitext(path)[1].lower()
        if extension == ".odt":
            body = self._load_odt(path)
        else:
            body = self._load_docx(path)
        self._browser.setHtml(
            "<html><body style='font-family: sans-serif;'>"
            f"{body}</body></html>"
        )

    def _load_docx(self, path: str) -> str:
        """@brief Convert a .docx file into a simplified HTML document.

        @details Refuses files over _MAX_DOC_BYTES and caps the number of
        rendered blocks to keep memory use bounded on hostile inputs.

        @param path: Absolute path of the DOCX file.
        @return HTML fragment; contains an error notice when parsing fails.
        """
        try:
            if os.path.getsize(path) > _MAX_DOC_BYTES:
                return "<p style='color:#a00'>Document is too large to display.</p>"
            import docx
            from docx.table import Table
            from docx.text.paragraph import Paragraph

            document = docx.Document(path)
        except ImportError:
            return "<p>python-docx package is not installed.</p>"
        except Exception as error:
            return f"<p style='color:#a00'>Could not open DOCX file: {html.escape(str(error))}</p>"

        parts: list[str] = []
        block_count = 0
        truncated = False
        try:
            blocks = list(document.iter_inner_content())
        except AttributeError:
            blocks = list(document.paragraphs) + list(document.tables)

        open_list = False
        for block in blocks:
            block_count += 1
            if block_count > _MAX_BLOCKS:
                truncated = True
                break
            if isinstance(block, Table):
                if open_list:
                    parts.append("</ul>")
                    open_list = False
                parts.append(self._table_to_html(block))
            elif isinstance(block, Paragraph):
                style = getattr(block.style, "name", "") or ""
                text = self._runs_to_html(block)
                if style.startswith("Heading"):
                    if open_list:
                        parts.append("</ul>")
                        open_list = False
                    try:
                        level = min(int(style.split()[-1]), 6)
                    except ValueError:
                        level = 2
                    parts.append(f"<h{level}>{text}</h{level}>")
                elif text:
                    if style.startswith("List"):
                        if not open_list:
                            parts.append("<ul>")
                            open_list = True
                        parts.append(f"<li>{text}</li>")
                    else:
                        if open_list:
                            parts.append("</ul>")
                            open_list = False
                        parts.append(f"<p>{text}</p>")
        if open_list:
            parts.append("</ul>")
        if truncated:
            parts.append(
                f"<p><b>[document truncated after {_MAX_BLOCKS} blocks]</b></p>"
            )
        return "".join(parts) or "<p>(empty document)</p>"

    def _runs_to_html(self, paragraph) -> str:
        """@brief Convert paragraph runs to HTML keeping bold/italic/underline.

        @param paragraph: docx.text.paragraph.Paragraph instance.
        @return HTML-escaped string with inline formatting tags.
        """
        pieces: list[str] = []
        for run in paragraph.runs:
            text = html.escape(run.text)
            if not text:
                continue
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            pieces.append(text)
        return "".join(pieces) or html.escape(paragraph.text)

    @staticmethod
    def _table_to_html(table) -> str:
        """@brief Convert a docx table to a bordered HTML table.

        @param table: docx.table.Table instance.
        @return HTML fragment with the first _MAX_TABLE_ROWS rows.
        """
        rows_html: list[str] = []
        for row_index, row in enumerate(islice(table.rows, _MAX_TABLE_ROWS)):
            tag = "th" if row_index == 0 else "td"
            cells = "".join(f"<{tag}>{html.escape(cell.text)}</{tag}>" for cell in row.cells)
            rows_html.append(f"<tr>{cells}</tr>")
        return (
            "<table border='1' cellspacing='0' cellpadding='4' "
            "style='border-collapse:collapse; margin:8px 0;'>"
            f"{''.join(rows_html)}</table>"
        )

    def _load_odt(self, path: str) -> str:
        """@brief Extract the text content of an OpenDocument .odt file.

        @details Reads content.xml straight from the ODT ZIP archive and
        converts text:h / text:p elements to headings and paragraphs. The
        uncompressed size of content.xml is capped to resist zip bombs.

        @param path: Absolute path of the ODT file.
        @return HTML fragment; contains an error notice when parsing fails.
        """
        try:
            with zipfile.ZipFile(path) as archive:
                info = archive.getinfo("content.xml")
                if info.file_size > _MAX_XML_BYTES:
                    return "<p style='color:#a00'>ODT content is too large to display.</p>"
                xml_bytes = archive.read("content.xml")
            root = ET.fromstring(xml_bytes)
        except (zipfile.BadZipFile, KeyError, ET.ParseError, OSError) as error:
            return f"<p style='color:#a00'>Could not read ODT file: {html.escape(str(error))}</p>"

        parts: list[str] = []
        block_count = 0
        for element in root.iter():
            if element.tag == f"{_ODT_TEXT_NS}h":
                block_count += 1
                if block_count > _MAX_BLOCKS:
                    parts.append(
                        f"<p><b>[document truncated after {_MAX_BLOCKS} blocks]</b></p>"
                    )
                    break
                try:
                    level = int(element.get(f"{_ODT_TEXT_NS}outline-level", "1"))
                except ValueError:
                    level = 1
                level = max(1, min(level, 6))
                parts.append(
                    f"<h{level}>{html.escape(''.join(element.itertext()))}</h{level}>"
                )
            elif element.tag == f"{_ODT_TEXT_NS}p":
                block_count += 1
                if block_count > _MAX_BLOCKS:
                    parts.append(
                        f"<p><b>[document truncated after {_MAX_BLOCKS} blocks]</b></p>"
                    )
                    break
                text = "".join(element.itertext())
                if text.strip():
                    parts.append(f"<p>{html.escape(text)}</p>")
        return "".join(parts) or "<p>(empty document)</p>"

    def cleanup(self) -> None:
        """@brief Clear the rendered document to free memory."""
        self._browser.clear()
